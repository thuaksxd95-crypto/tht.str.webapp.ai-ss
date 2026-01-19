import streamlit as st
import math
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import io
import xlsxwriter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 0. CẤU HÌNH & HÀM HỖ TRỢ
# ==========================================
st.set_page_config(
    page_title="Structure AI V21.1 (Stable Fix)",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Khởi tạo Session State
if 'current_floor_idx' not in st.session_state:
    st.session_state.current_floor_idx = 0

# CSS Styles
st.markdown("""
<style>
    .main-header { font-size:24px; font-weight: bold; color: #154360; border-bottom: 3px solid #2E86C1; padding-bottom: 8px; margin-bottom: 20px; text-transform: uppercase; }
    .sub-header { font-size:16px; font-weight: bold; color: #2C3E50; margin-top: 15px; margin-bottom: 5px; }
    div[data-testid="stExpander"] details summary p { font-weight: bold; font-size: 15px; }
</style>
""", unsafe_allow_html=True)

# --- REPORT ENGINE (Word không chứa ảnh để đảm bảo ổn định) ---
def create_docx_report(project_name, project_type, mat_info, load_info, design_results, mong_desc):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Header
    head = doc.add_heading('THUYẾT MINH TÍNH TOÁN KẾT CẤU SƠ BỘ', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"DỰ ÁN: {project_name.upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Loại công trình: {project_type} | Ngày lập: {pd.Timestamp.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("-" * 70)

    # 1. Tiêu chuẩn
    doc.add_heading('I. CÁC TIÊU CHUẨN ÁP DỤNG', level=1)
    stds = [
        "TCVN 2737:2023: Tải trọng và tác động - Tiêu chuẩn thiết kế.",
        "TCVN 5574:2018: Kết cấu bê tông và bê tông cốt thép - Tiêu chuẩn thiết kế.",
        "TCVN 9386:2012: Thiết kế công trình chịu động đất.",
        "TCVN 10304:2014: Móng cọc - Tiêu chuẩn thiết kế."
    ]
    for s in stds: doc.add_paragraph(s, style='List Bullet')

    # 2. Vật liệu
    doc.add_heading('II. THÔNG SỐ VẬT LIỆU', level=1)
    doc.add_paragraph(f"1. Bê tông: {mat_info['conc']} (Rb = {mat_info['rb']} MPa)")
    doc.add_paragraph(f"2. Cốt thép: {mat_info['steel']} (Rs = {mat_info['rs']} MPa)")
    doc.add_paragraph(f"3. Tải trọng sàn quy đổi: q = {load_info} kN/m2")

    # 3. Kết quả
    doc.add_heading('III. KẾT QUẢ TÍNH TOÁN & LỰA CHỌN', level=1)
    
    def add_df(df, title):
        doc.add_heading(title, level=2)
        if df.empty:
            doc.add_paragraph("Không áp dụng")
            return
        # Tạo bảng Word
        t = doc.add_table(df.shape[0]+1, df.shape[1])
        t.style = 'Table Grid'
        # Header
        for j, col in enumerate(df.columns): 
            t.cell(0, j).text = str(col)
        # Body
        for i, row in enumerate(df.itertuples(index=False)):
            for j, val in enumerate(row): 
                t.cell(i+1, j).text = str(val)
        doc.add_paragraph("") # Dòng trống

    add_df(design_results['San'], "1. Sàn (Slab)")
    add_df(design_results['Dam'], "2. Dầm (Beam)")
    add_df(design_results['Cot'], "3. Cột (Column)")
    if 'Vach' in design_results: add_df(design_results['Vach'], "4. Vách (Wall)")
    
    doc.add_heading('5. Móng (Foundation)', level=2)
    doc.add_paragraph(f"Phương án móng: {mong_desc}")
    add_df(design_results['Mong'], "Chi tiết móng:")

    doc.add_heading('IV. KẾT LUẬN', level=1)
    doc.add_paragraph("Phương án kết cấu sơ bộ đảm bảo khả năng chịu lực. Cần kiểm toán chi tiết trong giai đoạn TKKT.")
    
    # Lưu vào bộ nhớ đệm
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- CÁC HÀM HỖ TRỢ KHÁC ---
def color_status(val):
    color = 'red'
    if val == '✅ ĐẠT': color = 'green'
    elif val == '⚠️ DƯ': color = '#B7950B'
    return f'color: {color}; font-weight: bold'

def parse_input_string(input_str):
    try:
        items = input_str.split(',')
        result = []
        for item in items:
            item = item.strip().lower()
            if 'x' in item:
                val, count = item.split('x')
                result.extend([float(val)] * int(count))
            else:
                if item: result.append(float(item))
        return result
    except:
        return []

def get_material_properties(grade_conc, grade_steel):
    rb_map = {"B20": 11.5, "B25": 14.5, "B30": 17.0, "B35": 19.5, "B40": 22.0, "B45": 25.0}
    rs_map = {"CB240-T": 210, "CB300-V": 260, "CB400-V": 350, "CB500-V": 435}
    return rb_map.get(grade_conc, 14.5), rs_map.get(grade_steel, 350)

def to_excel(dfs):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        for sheet_name, df in dfs.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            worksheet = writer.sheets[sheet_name]
            for i, col in enumerate(df.columns):
                col_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.set_column(i, i, col_len)
    return output.getvalue()

# Dữ liệu chuẩn
RB_MAP = {"B15": 8.5, "B20": 11.5, "B25": 14.5, "B30": 17.0, "B35": 19.5, "B40": 22.0, "B45": 25.0, "B50": 27.5}
RS_MAP = {"CB240-T": 210, "CB300-T": 260, "CB300-V": 260, "CB400-V": 350, "CB500-V": 435, "CB600-V": 520}
Q_DEFAULTS = {"Nhà phố/Biệt thự": 10.0, "Văn phòng/Khách sạn": 14.0, "Chung cư cao tầng": 14.5}

# ==========================================
# 1. SIDEBAR INPUT
# ==========================================
with st.sidebar:
    st.image("https://img.icons8.com/color/96/000000/structural.png", width=60)
    st.title("THIẾT LẬP DỰ ÁN")
    
    with st.expander("1. Thông Tin Chung", expanded=True):
        project_name = st.text_input("Tên dự án", "Tòa nhà Văn phòng A")
        pt_opts = list(Q_DEFAULTS.keys()) + ["Tùy chỉnh..."]
        pt_sel = st.selectbox("Loại công trình", pt_opts, index=1)
        if pt_sel == "Tùy chỉnh...":
            project_type = st.text_input("Nhập tên loại CT", "Nhà xưởng/Kho")
            default_q = 10.0
        else:
            project_type = pt_sel
            default_q = Q_DEFAULTS[pt_sel]
        has_shearwall = st.checkbox("Có sử dụng Vách cứng?", value=False)

    with st.expander("2. Hệ Lưới & Cao Độ", expanded=True):
        st.markdown("**A. Cao độ tầng (m):**")
        height_mode = st.radio("Chế độ nhập:", ["Điển hình", "Chi tiết"], horizontal=True, label_visibility="collapsed")
        if height_mode == "Điển hình":
            num_floors = st.number_input("Tổng số tầng", 1, 100, 10)
            h_typ = st.number_input("Chiều cao điển hình (m)", 2.0, 6.0, 3.3)
            floor_heights = [h_typ] * num_floors
        else:
            h_str = st.text_input("Nhập chuỗi (VD: 4.5, 3.3x9)", "4.5, 3.3x9")
            floor_heights = parse_input_string(h_str)
            num_floors = len(floor_heights)
            st.caption(f"Tổng: {num_floors} tầng | Cao: {sum(floor_heights):.1f}m")
        st.markdown("---")
        st.markdown("**B. Lưới trục (m):**")
        grid_x_str = st.text_input("Khoảng cách trục X", "6, 7, 6")
        grid_y_str = st.text_input("Khoảng cách trục Y", "5, 5, 5")
        lx_list = parse_input_string(grid_x_str)
        ly_list = parse_input_string(grid_y_str)
        l_max = max(max(lx_list, default=0), max(ly_list, default=0))
        l_min = min(max(lx_list, default=0), max(ly_list, default=0))
        area_trib = max(lx_list, default=0) * max(ly_list, default=0)

    with st.expander("3. Thông Số Vật Liệu", expanded=False):
        conc_opts = list(RB_MAP.keys()) + ["Tùy chỉnh..."]
        conc_sel = st.selectbox("Bê tông (Concrete)", conc_opts, index=3)
        if conc_sel == "Tùy chỉnh...":
            rb = st.number_input("Nhập Rb (MPa)", 1.0, 100.0, 14.5, step=0.5)
            conc_grade = "Custom"
        else:
            rb = RB_MAP[conc_sel]
            conc_grade = conc_sel
            st.caption(f"Rb = {rb} MPa")
        
        st.markdown("---")
        main_steel_opts = list(RS_MAP.keys()) + ["Tùy chỉnh..."]
        main_steel_sel = st.selectbox("Thép chủ (Main Bar)", main_steel_opts, index=3)
        if main_steel_sel == "Tùy chỉnh...":
            rs = st.number_input("Nhập Rs (MPa)", 100, 1000, 350, step=10)
            steel_main = "Custom"
        else:
            rs = RS_MAP[main_steel_sel]
            steel_main = main_steel_sel
            st.caption(f"Rs = {rs} MPa")
            
        stir_opts = list(RS_MAP.keys()) + ["Tùy chỉnh..."]
        stir_sel = st.selectbox("Thép đai (Stirrup)", stir_opts, index=0)
        if stir_sel == "Tùy chỉnh...":
            rsw = st.number_input("Nhập Rsw (MPa)", 100, 1000, 170, step=10)
            steel_stirrup = "Custom"
        else:
            rsw = RS_MAP[stir_sel]
            steel_stirrup = stir_sel
        
        st.markdown("---")
        q_load = st.number_input("Tải trọng sàn (kN/m2)", value=default_q)

    with st.expander("4. Cấu Kiện Cột", expanded=False):
        col_shape = st.radio("Hình dạng:", ["Chữ nhật", "Vuông"], horizontal=True)
        col_orient = st.radio("Phương cột CN:", ["Dọc nhà (Theo Y)", "Ngang nhà (Theo X)"], index=0)
        b_col_fixed = st.number_input("Cạnh b cố định (mm)", 150, 1000, 220, step=10)
        k_safety = 1.15

    with st.expander("5. Cấu Kiện Móng", expanded=False):
        found_type = st.selectbox("Loại móng", ["Móng Cọc (Pile)", "Móng Đơn/Băng"])
        if found_type == "Móng Cọc (Pile)":
            pile_std = ["Vuông 200x200", "Vuông 250x250", "Vuông 300x300", "Vuông 350x350", "Vuông 400x400", "Ly tâm D300", "Ly tâm D350", "Ly tâm D400", "Ly tâm D500", "Ly tâm D600", "Khoan nhồi D800", "Khoan nhồi D1000", "Tùy chỉnh..."]
            pile_sel = st.selectbox("Chọn loại cọc", pile_std, index=2)
            if pile_sel == "Tùy chỉnh...":
                d_pile = st.number_input("Kích thước/Đường kính cọc (mm)", 100, 2000, 400)
                pile_type = f"D{d_pile} (Custom)"
            else:
                pile_type = pile_sel
                if "Vuông" in pile_sel: d_pile = int(pile_sel.split(' ')[1].split('x')[0])
                elif "D" in pile_sel: d_pile = int(pile_sel.split('D')[1])
                else: d_pile = 400
            p_pile = st.number_input("Sức chịu tải TK (Tấn)", 10, 2000, 45)
        else:
            r_dat = st.number_input("Cường độ đất nền R (kg/cm2)", 0.5, 10.0, 1.5)

# ==========================================
# 2. CALCULATION ENGINE
# ==========================================
# Sàn
hs_calc = (l_min * 1000) / 35
hs_select = max(100, math.ceil(hs_calc / 10) * 10)
df_slab = pd.DataFrame([{"Cấu kiện": "Sàn điển hình", "Hoạt tải (kN/m2)": q_load, "Nhịp ngắn L (m)": l_min, "Công thức": "L/35", "Chiều dày YC (mm)": hs_calc, "Chiều dày CHỌN (mm)": int(hs_select), "Hệ số AT": hs_select/hs_calc if hs_calc else 0, "Trạng thái": "✅ ĐẠT" if hs_select >= hs_calc else "⛔ KHÔNG ĐẠT"}])

# Dầm
hd_calc = (l_max * 1000) / 12; hd_select = math.ceil(hd_calc / 50) * 50
bd_calc = 0.4 * hd_select; bd_select = max(200, math.ceil(bd_calc / 50) * 50)
if hd_select >= 700 and bd_select < 300: bd_select = 300
hd_sec = (l_max * 1000) / 16; hd_sec_s = math.ceil(hd_sec / 50) * 50
bd_sec_s = max(200, math.ceil(0.4 * hd_sec_s / 50) * 50)
df_beam = pd.DataFrame([
    {"Cấu kiện": "Dầm khung chính", "Nhịp lớn L (m)": l_max, "Công thức": "L/12", "Chiều cao YC (mm)": hd_calc, "Tiết diện CHỌN (mm)": f"{int(bd_select)}x{int(hd_select)}", "Hệ số AT": hd_select/hd_calc if hd_calc else 0, "Trạng thái": "✅ ĐẠT"},
    {"Cấu kiện": "Dầm phụ", "Nhịp lớn L (m)": l_max, "Công thức": "L/16", "Chiều cao YC (mm)": hd_sec, "Tiết diện CHỌN (mm)": f"{int(bd_sec_s)}x{int(hd_sec_s)}", "Hệ số AT": hd_sec_s/hd_sec if hd_sec else 0, "Trạng thái": "✅ ĐẠT"}
])

# Cột - Fix lỗi NULL bằng vòng lặp chuẩn
col_schedule = []
floors = list(range(1, num_floors + 1))[::-1]
group_map = {}
for f in floors:
    idx = (f - 1) // 3
    if idx not in group_map: group_map[idx] = []
    group_map[idx].append(f)

for grp_id in sorted(group_map.keys(), reverse=True):
    floor_list = group_map[grp_id]
    n_supported = num_floors - min(floor_list) + 1
    N_calc = k_safety * q_load * area_trib * n_supported
    Ac_req = (N_calc * 1000) / rb
    if col_shape == "Vuông":
        side = math.sqrt(Ac_req); b_sel = h_sel = math.ceil(side / 50) * 50
    else:
        h_req = Ac_req / b_col_fixed; h_sel = math.ceil(h_req / 50) * 50; b_sel = b_col_fixed
    if h_sel < 200: h_sel = 200
    if b_sel < 200: b_sel = 200
    status = "✅ ĐẠT" if (b_sel * h_sel) >= Ac_req else "⛔ KHÔNG ĐẠT"
    col_schedule.append({"Vị trí": f"Tầng {min(floor_list)}-{max(floor_list)}", "Tải N (kN)": N_calc, "A_yc (cm2)": Ac_req/100, "Tiết diện": f"{int(b_sel)}x{int(h_sel)}", "A_chon (cm2)": int(b_sel*h_sel/100), "Ratio": (b_sel*h_sel)/Ac_req if Ac_req else 0, "Trạng thái": status})
df_col = pd.DataFrame(col_schedule).iloc[::-1].reset_index(drop=True)

# Vách
df_wall = pd.DataFrame()
if has_shearwall:
    h_max = max(floor_heights) if floor_heights else 3.3
    tw_calc = h_max * 1000 / 20; tw_select = max(200, math.ceil(tw_calc / 50) * 50)
    df_wall = pd.DataFrame([{"Cấu kiện": "Vách cứng điển hình", "Chiều cao tầng H (m)": h_max, "Công thức": "H/20", "Chiều dày YC (mm)": tw_calc, "Chiều dày CHỌN (mm)": int(tw_select), "Hệ số AT": tw_select/tw_calc if tw_calc else 0, "Trạng thái": "✅ ĐẠT"}])

# Móng
N_footing = df_col.iloc[-1]["Tải N (kN)"] * 1.1 if not df_col.empty else 0
mong_desc = ""; mong_detail = ""
if found_type == "Móng Cọc (Pile)":
    n_pile_calc = N_footing / (p_pile * 9.81); n_pile = math.ceil(n_pile_calc * 1.2)
    spacing = 3 * (d_pile/1000); edge = 0.7 * (d_pile/1000)
    w = l = round(math.sqrt(n_pile * spacing**2), 1) if n_pile > 4 else round(spacing + d_pile/1000 + 2*edge, 2)
    mong_desc = f"{n_pile} cọc {pile_type}"; mong_detail = f"Đài {w}x{l}m (P={p_pile}T)"
    df_found = pd.DataFrame([{"Cấu kiện": f"Móng ({found_type})", "Tải chân cột N (kN)": N_footing, "Sức chịu tải P (T)": p_pile, "Số cọc YC": n_pile_calc, "Số cọc CHỌN": int(n_pile), "Kích thước / Ghi chú": mong_detail, "Trạng thái": "✅ ĐẠT"}])
else:
    R_convert = r_dat * 100; F_req = N_footing / (R_convert - 20); side = math.ceil(math.sqrt(F_req)*10)/10
    mong_desc = f"Móng đơn B={side}m"; mong_detail = f"R={r_dat}kg/cm2"
    df_found = pd.DataFrame([{"Cấu kiện": f"Móng ({found_type})", "Tải chân cột N (kN)": N_footing, "R đất (kg/cm2)": r_dat, "Diện tích YC (m2)": F_req, "Diện tích CHỌN (m2)": side*side, "Kích thước / Ghi chú": mong_detail, "Trạng thái": "✅ ĐẠT"}])

data_collection = {"San": df_slab, "Dam": df_beam, "Cot": df_col, "Mong": df_found}
if not df_wall.empty: data_collection["Vach"] = df_wall

# ==========================================
# 3. GLOBAL GRAPHICS GENERATION
# ==========================================
cum_x = [0]; grid_labels_x = ["1"]
for i, val in enumerate(lx_list): cum_x.append(cum_x[-1] + val); grid_labels_x.append(str(i + 2))
cum_y = [0]; grid_labels_y = ["A"]
for i, val in enumerate(ly_list): cum_y.append(cum_y[-1] + val); grid_labels_y.append(chr(65 + i + 1))
cum_z = [0]; level_labels = ["Móng"]
for i, val in enumerate(floor_heights): cum_z.append(cum_z[-1] + val); level_labels.append(f"Tầng {i+1}" if i < len(floor_heights)-1 else "Mái")

if st.session_state.current_floor_idx >= len(floor_heights): st.session_state.current_floor_idx = len(floor_heights) - 1
if st.session_state.current_floor_idx < 0: st.session_state.current_floor_idx = 0
current_z = cum_z[st.session_state.current_floor_idx + 1]
current_label = level_labels[st.session_state.current_floor_idx + 1]

if not df_col.empty:
    c_dim = df_col.iloc[0]["Tiết diện"].split('x')
    dim1 = float(c_dim[0])/1000; dim2 = float(c_dim[1])/1000
    if col_shape == "Vuông": bc_m, hc_m = dim1, dim1
    else:
        if col_orient == "Ngang nhà (Theo X)": bc_m, hc_m = max(dim1, dim2), min(dim1, dim2)
        else: bc_m, hc_m = min(dim1, dim2), max(dim1, dim2)
else: bc_m = hc_m = 0.2

# 3.1 DRAW PLAN
fig_plan = go.Figure()
for x, label in zip(cum_x, grid_labels_x):
    fig_plan.add_trace(go.Scatter(x=[x, x], y=[min(cum_y)-1, max(cum_y)+1], mode='lines+text', line=dict(color='#BDC3C7', width=1, dash='dash'), text=[None, label], textposition="top center", hoverinfo='skip'))
    fig_plan.add_trace(go.Scatter(x=[x], y=[min(cum_y)-1], mode='markers+text', marker=dict(size=25, color='white', line=dict(color='black', width=1)), text=label, textposition="middle center", showlegend=False, hoverinfo='skip'))
for y, label in zip(cum_y, grid_labels_y):
    fig_plan.add_trace(go.Scatter(x=[min(cum_x)-1, max(cum_x)+1], y=[y, y], mode='lines+text', line=dict(color='#BDC3C7', width=1, dash='dash'), text=[None, label], textposition="middle right", hoverinfo='skip'))
    fig_plan.add_trace(go.Scatter(x=[min(cum_x)-1], y=[y], mode='markers+text', marker=dict(size=25, color='white', line=dict(color='black', width=1)), text=label, textposition="middle center", showlegend=False, hoverinfo='skip'))
bx, by = [], []
for y in cum_y: bx.extend([min(cum_x), max(cum_x), None]); by.extend([y, y, None])
for x in cum_x: bx.extend([x, x, None]); by.extend([min(cum_y), max(cum_y), None])
fig_plan.add_trace(go.Scatter(x=bx, y=by, mode='lines', line=dict(color='#2980B9', width=3), name='Dầm', hoverinfo='text'))
shapes = []
for x in cum_x:
    for y in cum_y: shapes.append(dict(type="rect", x0=x-bc_m/2, y0=y-hc_m/2, x1=x+bc_m/2, y1=y+hc_m/2, fillcolor="#E74C3C", line=dict(width=0)))
fig_plan.update_layout(shapes=shapes, xaxis=dict(visible=False, fixedrange=False, range=[min(cum_x)-2, max(cum_x)+2]), yaxis=dict(visible=False, scaleanchor="x", fixedrange=False, range=[min(cum_y)-2, max(cum_y)+2]), margin=dict(l=10,r=10,t=10,b=10), height=500, dragmode="pan", showlegend=False, title="MẶT BẰNG KẾT CẤU")

# 3.2 DRAW ELEVATION
fig_elev = go.Figure()
x_min, x_max = min(cum_x) - 1, max(cum_x) + 1
for i, z in enumerate(cum_z):
    label = level_labels[i]
    line_color = '#7F8C8D'; line_width = 1
    fig_elev.add_trace(go.Scatter(x=[x_min, x_max + 1.5], y=[z, z], mode='lines', line=dict(color=line_color, width=line_width, dash='dot'), hoverinfo='skip'))
    marker_x = x_max + 1.5
    fig_elev.add_trace(go.Scatter(x=[marker_x], y=[z], mode='markers', marker=dict(symbol='triangle-down', size=15, color=line_color, line=dict(width=1, color=line_color)), hoverinfo='skip', showlegend=False))
    fig_elev.add_trace(go.Scatter(x=[marker_x], y=[z + 0.2], mode='text', text=[f"{label} (+{z:.2f})"], textposition="top center", textfont=dict(color=line_color, size=12), hoverinfo='skip', showlegend=False))
for x, label in zip(cum_x, grid_labels_x):
    fig_elev.add_trace(go.Scatter(x=[x, x], y=[-1, max(cum_z)+1], mode='lines', line=dict(color='#BDC3C7', width=1, dash='dash'), showlegend=False, hoverinfo='skip'))
    fig_elev.add_trace(go.Scatter(x=[x], y=[-1.5], mode='markers+text', marker=dict(size=25, color='white', line=dict(color='black', width=1)), text=label, textposition="middle center", showlegend=False))
shapes_elev = []
for x in cum_x: shapes_elev.append(dict(type="rect", x0=x-bc_m/2, y0=0, x1=x+bc_m/2, y1=max(cum_z), fillcolor="#BDC3C7", opacity=0.5, line=dict(width=0)))
for z in cum_z[1:]:
    for j in range(len(cum_x)-1): shapes_elev.append(dict(type="rect", x0=cum_x[j], y0=z-0.5, x1=cum_x[j+1], y1=z, fillcolor="#3498DB", opacity=0.5, line=dict(width=0)))
fig_elev.update_layout(shapes=shapes_elev, xaxis=dict(visible=False, fixedrange=False, range=[x_min-1, x_max+4]), yaxis=dict(visible=False, scaleanchor="x", fixedrange=False, range=[-2, max(cum_z)+2]), margin=dict(l=10,r=10,t=10,b=10), height=500, dragmode="pan", showlegend=False, title="MẶT ĐỨNG KHUNG")

# ==========================================
# 4. MAIN APP UI
# ==========================================
st.title(f"📐 {project_name.upper()}")
st.markdown(f"**Loại:** {project_type} | **Vật liệu:** BT {conc_grade} (Rb={rb}), Thép {steel_main} (Rs={rs})")

tab1, tab2 = st.tabs(["📊 BẢN VẼ & BẢNG TÍNH", "📝 THUYẾT MINH"])

with tab1:
    c_nav1, c_nav2, c_nav3 = st.columns([1, 2, 1])
    with c_nav2:
        col_btn1, col_info, col_btn2 = st.columns([1, 2, 1])
        if col_btn1.button("⬇️ Xuống tầng", use_container_width=True):
            if st.session_state.current_floor_idx > 0: st.session_state.current_floor_idx -= 1; st.rerun()
        col_info.markdown(f"<div style='text-align:center; font-weight:bold; font-size:18px; color:#2E86C1; border: 1px solid #ddd; padding: 5px; border-radius: 5px;'>{current_label} (+{current_z:.2f}m)</div>", unsafe_allow_html=True)
        if col_btn2.button("⬆️ Lên tầng", use_container_width=True):
            if st.session_state.current_floor_idx < len(floor_heights) - 1: st.session_state.current_floor_idx += 1; st.rerun()

    col_plan, col_elev = st.columns([1, 1])
    with col_plan: st.plotly_chart(fig_plan, use_container_width=True, config={'scrollZoom': True, 'displayModeBar': True})
    with col_elev: st.plotly_chart(fig_elev, use_container_width=True, config={'scrollZoom': True, 'displayModeBar': True})

    st.markdown("---")
    
    col_h1, col_h2 = st.columns([3, 1])
    with col_h1: st.markdown('<p class="main-header">BẢNG TỔNG HỢP & KIỂM TRA KẾT QUẢ TÍNH TOÁN</p>', unsafe_allow_html=True)
    with col_h2: st.download_button("📥 Xuất Excel", data=to_excel(data_collection), file_name=f"{project_name}_Calc.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.markdown('<p class="sub-header">🟦 1. KẾT CẤU BẢN SÀN (SLAB CHECK)</p>', unsafe_allow_html=True)
    st.dataframe(df_slab.style.map(color_status, subset=['Trạng thái']), use_container_width=True, hide_index=True, column_config={"Hoạt tải (kN/m2)": st.column_config.NumberColumn(format="%.2f"), "Nhịp ngắn L (m)": st.column_config.NumberColumn(format="%.2f"), "Chiều dày YC (mm)": st.column_config.NumberColumn(format="%.2f"), "Chiều dày CHỌN (mm)": st.column_config.NumberColumn(format="%d"), "Hệ số AT": st.column_config.NumberColumn(format="%.2f")})

    st.markdown('<p class="sub-header">🟩 2. KẾT CẤU DẦM KHUNG (BEAM CHECK)</p>', unsafe_allow_html=True)
    st.dataframe(df_beam.style.map(color_status, subset=['Trạng thái']), use_container_width=True, hide_index=True, column_config={"Nhịp lớn L (m)": st.column_config.NumberColumn(format="%.2f"), "Chiều cao YC (mm)": st.column_config.NumberColumn(format="%.2f"), "Hệ số AT": st.column_config.NumberColumn(format="%.2f")})

    st.markdown('<p class="sub-header">🟥 3. KẾT CẤU CỘT (COLUMN SCHEDULE & CHECK)</p>', unsafe_allow_html=True)
    st.dataframe(df_col.style.map(color_status, subset=['Trạng thái']), use_container_width=True, hide_index=True, column_config={"Vị trí": st.column_config.TextColumn("Zone Tầng", width="small"), "Tải N (kN)": st.column_config.ProgressColumn("Lực Dọc N (kN)", format="%.2f", min_value=0, max_value=int(df_col["Tải N (kN)"].max()*1.1)), "A_yc (cm2)": st.column_config.NumberColumn("Diện tích YC", format="%.2f"), "A_chon (cm2)": st.column_config.NumberColumn("Diện tích CHỌN", format="%d"), "Ratio": st.column_config.NumberColumn("HS An Toàn", format="%.2f")})

    if has_shearwall and not df_wall.empty:
        st.markdown('<p class="sub-header">🟧 4. KẾT CẤU VÁCH CỨNG (SHEAR WALL)</p>', unsafe_allow_html=True)
        st.dataframe(df_wall.style.map(color_status, subset=['Trạng thái']), use_container_width=True, hide_index=True, column_config={"Chiều cao tầng H (m)": st.column_config.NumberColumn(format="%.2f"), "Chiều dày YC (mm)": st.column_config.NumberColumn(format="%.2f"), "Chiều dày CHỌN (mm)": st.column_config.NumberColumn(format="%d"), "Hệ số AT": st.column_config.NumberColumn(format="%.2f")})

    st.markdown('<p class="sub-header">🟫 5. KẾT CẤU MÓNG (FOUNDATION CHECK)</p>', unsafe_allow_html=True)
    mong_fmt = {"Tải chân cột N (kN)": st.column_config.NumberColumn(format="%.2f"), "Sức chịu tải P (T)": st.column_config.NumberColumn(format="%.2f"), "Số cọc YC": st.column_config.NumberColumn(format="%.2f"), "Số cọc CHỌN": st.column_config.NumberColumn(format="%d")}
    if found_type != "Móng Cọc (Pile)": mong_fmt = {"Tải chân cột N (kN)": st.column_config.NumberColumn(format="%.2f"), "R đất (kg/cm2)": st.column_config.NumberColumn(format="%.2f"), "Diện tích YC (m2)": st.column_config.NumberColumn(format="%.2f"), "Diện tích CHỌN (m2)": st.column_config.NumberColumn(format="%.2f")}
    st.dataframe(df_found.style.map(color_status, subset=['Trạng thái']), use_container_width=True, hide_index=True, column_config=mong_fmt)

with tab2:
    st.markdown('<p class="report-title">THUYẾT MINH TÍNH TOÁN KẾT CẤU</p>', unsafe_allow_html=True)
    st.markdown(f'<p class="report-sub">Dự án: {project_name} | Ngày lập: {pd.Timestamp.now().strftime("%d/%m/%Y")}</p>', unsafe_allow_html=True)
    st.markdown("---")
    st.header("I. CÁC TIÊU CHUẨN ÁP DỤNG")
    st.markdown("""
    * **TCVN 2737:2023:** Tải trọng và tác động.
    * **TCVN 5574:2018:** Kết cấu bê tông và bê tông cốt thép.
    * **TCVN 9386:2012:** Thiết kế công trình chịu động đất.
    * **TCVN 10304:2014:** Móng cọc - Tiêu chuẩn thiết kế.
    """)
    st.header("II. THÔNG SỐ ĐẦU VÀO")
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("1. Vật Liệu")
        st.markdown(f"- **Bê tông:** {conc_grade} ($R_b = {rb}$ MPa)")
        st.markdown(f"- **Thép chủ:** {steel_main} ($R_s = {rs}$ MPa)")
        st.markdown(f"- **Thép đai:** {steel_stirrup}")
    with c2:
        st.subheader("2. Tải Trọng")
        st.markdown(f"- **Tải sàn quy đổi:** $q = {q_load}$ $kN/m^2$")
        st.caption("(Trọng lượng bản thân + hoàn thiện + tường + hoạt tải)")
    
    st.header("III. KẾT QUẢ TÍNH TOÁN SƠ BỘ")
    st.subheader("1. Bản Sàn (Slab)")
    st.markdown(f"Chiều dày sàn chọn sơ bộ: $h_s = D/m \\cdot L$")
    st.latex(r"h_{yc} = \frac{L_{min}}{35} = " + f"{hs_calc:.1f} mm")
    st.success(f"👉 **CHỌN: Chiều dày sàn {int(hs_select)} mm**")
    
    st.subheader("2. Dầm Khung (Beam)")
    st.markdown(f"Chiều cao dầm chính: $h_d = (1/8 \div 1/12)L$")
    st.latex(r"h_{yc} \approx \frac{L_{max}}{12} = " + f"{hd_calc:.0f} mm")
    st.success(f"👉 **CHỌN: Dầm {int(bd_select)}x{int(hd_select)} mm**")
    
    st.subheader("3. Cột (Column)")
    st.markdown("Diện tích tiết diện cột sơ bộ:")
    st.latex(r"A_{yc} = \frac{k \cdot N}{R_b}")
    
    st.subheader("4. Móng (Foundation)")
    if found_type == "Móng Cọc (Pile)":
        st.markdown("Số lượng cọc sơ bộ:")
        st.latex(r"n = \frac{1.2 \cdot N_{chan}}{P_{tk}}")
        st.success(f"👉 **CHỌN: {mong_desc}**")
    else:
        st.markdown("Diện tích đáy móng nông:")
        st.latex(r"F = \frac{N_{chan}}{R_{dat} - \gamma H}")
        st.success(f"👉 **CHỌN: {mong_desc}**")

    st.markdown("---")
    # Export Docx
    mat_info = {'conc': conc_grade, 'rb': rb, 'steel': steel_main, 'rs': rs}
    docx_file = create_docx_report(project_name, project_type, mat_info, q_load, data_collection, mong_desc)
    
    st.download_button(
        label="📄 Tải Thuyết Minh (.docx)",
        data=docx_file,
        file_name=f"{project_name}_ThuyetMinh.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
